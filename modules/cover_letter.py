import io
import re
import streamlit as st
from config import ANTHROPIC_API_KEY, CLAUDE_FAST


def generate_cover_letter(profile: dict, job: dict, language: str = "en", tone: str = "professional") -> str:
    if not ANTHROPIC_API_KEY or ANTHROPIC_API_KEY == "your_key_here":
        return _fallback_letter(profile, job, language)

    try:
        import anthropic
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    except ImportError:
        return _fallback_letter(profile, job, language)

    name = profile.get("name", "")
    exp = profile.get("experience_text", "")[:800]
    skills = ", ".join(profile.get("skills", [])[:12])
    summary = profile.get("summary", "")
    role = job.get("title", "")
    company = job.get("company", "")
    jd = job.get("description", "")[:1000]
    lang_instruction = "Write in Hebrew (RTL)" if language == "he" else "Write in English"

    prompt = f"""Write a tailored cover letter for {name} applying to {role} at {company}.

{lang_instruction}
Tone: {tone}
Length: 3 focused paragraphs + closing

Candidate background:
Summary: {summary}
Experience: {exp}
Skills: {skills}

Job description excerpt:
{jd}

Rules:
- Do NOT exaggerate or invent experience
- Reference specific skills that match the JD naturally
- Sound human, not AI-generated
- Include a specific reason why this company is attractive
- End with a clear call to action

Write the letter only, no extra commentary."""

    try:
        msg = client.messages.create(
            model=CLAUDE_FAST,
            max_tokens=900,
            messages=[{"role": "user", "content": prompt}],
        )
        return msg.content[0].text.strip()
    except Exception as e:
        return _fallback_letter(profile, job, language) + f"\n\n[Note: Claude error — {e}]"


def _fallback_letter(profile: dict, job: dict, language: str) -> str:
    name = profile.get("name", "")
    role = job.get("title", "")
    company = job.get("company", "")
    skills = ", ".join(profile.get("skills", [])[:5])

    if language == "he":
        return f"""שלום,

אני {name}, ואני מגיש מועמדות לתפקיד {role} ב-{company}.

לאחר בחינת תיאור התפקיד, אני רואה התאמה גבוהה בין הניסיון שצברתי לבין הדרישות.
הכישורים המרכזיים שאני מביא: {skills}.

בניסיוני הקודם פיתחתי יכולות חזקות בניתוח נתונים, בניית דשבורדים ועבודה עם stakeholders מגוונים.
אשמח לשוחח על הזדמנות ההצטרפות לצוות.

בברכה,
{name}"""
    else:
        return f"""Dear Hiring Team,

I am {name}, writing to express my interest in the {role} position at {company}.

After reviewing the job description, I see a strong alignment between my experience and your requirements.
Key skills I bring: {skills}.

In my previous roles, I developed strong capabilities in data analysis, dashboard development, and cross-functional collaboration.
I look forward to discussing how I can contribute to your team.

Best regards,
{name}"""


def export_letter_docx(text: str, profile: dict) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return b""

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    for para_text in text.split("\n"):
        p = doc.add_paragraph(para_text)
        # add_paragraph("") creates no runs — indexing p.runs[0] would crash
        # on every blank line between paragraphs.
        for r in p.runs:
            r.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render(lang: str):
    profile = st.session_state.profile
    selected_job = st.session_state.get("selected_job", {})

    title = "✉️ מכתב מקדים" if lang == "he" else "✉️ Cover Letter"
    st.markdown(f'<div class="section-title">{title}</div>', unsafe_allow_html=True)

    col_l, col_r = st.columns([1, 1])

    with col_l:
        # Job info
        st.markdown("**" + ("פרטי המשרה" if lang == "he" else "Job Details") + "**")
        job_title = st.text_input(
            "תפקיד" if lang == "he" else "Role",
            value=selected_job.get("title", "Data Analyst"),
        )
        job_company = st.text_input(
            "חברה" if lang == "he" else "Company",
            value=selected_job.get("company", ""),
        )
        job_desc = st.text_area(
            "תיאור משרה (אופציונלי)" if lang == "he" else "Job Description (optional)",
            value=selected_job.get("description", "")[:500],
            height=120,
        )

        # Options
        letter_lang = st.radio(
            "שפת המכתב" if lang == "he" else "Letter Language",
            ["English", "עברית"],
            horizontal=True,
        )
        letter_lang_code = "he" if letter_lang == "עברית" else "en"

        tone = st.select_slider(
            "טון" if lang == "he" else "Tone",
            options=["Formal", "Professional", "Friendly", "Enthusiastic"],
            value="Professional",
        )

        has_claude = ANTHROPIC_API_KEY and ANTHROPIC_API_KEY != "your_key_here"
        gen_label = "✨ " + ("צור מכתב" if lang == "he" else "Generate Letter")
        if st.button(gen_label, type="primary"):
            job_dict = {
                "title": job_title,
                "company": job_company,
                "description": job_desc,
            }
            with st.spinner("כותב מכתב..." if lang == "he" else "Writing letter..."):
                letter = generate_cover_letter(profile, job_dict, letter_lang_code, tone.lower())
                st.session_state.cover_letter_text = letter

    with col_r:
        preview_title = "המכתב" if lang == "he" else "Cover Letter"
        st.markdown(f"**{preview_title}**")

        letter_text = st.session_state.get("cover_letter_text", "")
        rtl = "rtl" if letter_lang_code == "he" else "ltr"

        edited_letter = st.text_area(
            preview_title,
            value=letter_text,
            height=380,
            label_visibility="collapsed",
            key="letter_editor",
        )
        st.session_state.cover_letter_text = edited_letter

    # ── Export ────────────────────────────────────────────────────────────────
    if edited_letter:
        st.markdown("---")
        c1, c2, c3 = st.columns(3)
        with c1:
            docx_bytes = export_letter_docx(edited_letter, profile)
            if docx_bytes:
                st.download_button(
                    "📥 .DOCX",
                    data=docx_bytes,
                    file_name="Cover_Letter.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
        with c2:
            st.download_button(
                "📥 .TXT",
                data=edited_letter.encode("utf-8"),
                file_name="Cover_Letter.txt",
                mime="text/plain",
                use_container_width=True,
            )
        with c3:
            copy_label = "📋 " + ("העתק ללוח" if lang == "he" else "Copy to Clipboard")
            if st.button(copy_label, use_container_width=True):
                st.code(edited_letter, language=None)
                st.toast("📋 " + ("הועתק!" if lang == "he" else "Copied to code block above!"))

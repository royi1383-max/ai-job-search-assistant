import json
import os
import re
import streamlit as st
from config import PROFILE_PATH, DATA_DIR, CV_FOLDER, TARGET_ROLES


# ── Persistence ──────────────────────────────────────────────────────────────

def save_profile(data: dict):
    os.makedirs(DATA_DIR, exist_ok=True)
    with open(PROFILE_PATH, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    st.session_state.profile = data


def load_profile() -> dict:
    if os.path.exists(PROFILE_PATH):
        with open(PROFILE_PATH, encoding="utf-8") as f:
            return json.load(f)
    return {}


def get_skills_list(profile: dict) -> list[str]:
    return profile.get("skills", [])


# ── CV parsing ───────────────────────────────────────────────────────────────

def parse_cv_pdf(path: str) -> dict:
    """Extract structured data from a PDF CV using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return {}
    try:
        doc = fitz.open(path)
        text = "\n".join(page.get_text() for page in doc)
        doc.close()
    except Exception:
        return {}
    return _extract_fields(text)


def parse_cv_docx(path: str) -> dict:
    """Extract structured data from a DOCX CV using python-docx."""
    try:
        from docx import Document
    except ImportError:
        return {}
    try:
        doc = Document(path)
        paragraphs = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                paragraphs.append(t)
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t and t not in paragraphs:
                        paragraphs.append(t)
        text = "\n".join(paragraphs)
    except Exception:
        return {}
    return _extract_fields(text)


def parse_cv_file(path: str) -> dict:
    """Route to PDF or DOCX parser based on extension."""
    if path.lower().endswith(".docx"):
        return parse_cv_docx(path)
    return parse_cv_pdf(path)


def _extract_fields(text: str) -> dict:
    data = {}

    # Name — first non-empty line that looks like a name
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    for line in lines[:5]:
        if re.match(r'^[A-Za-z֐-׿][\w\s֐-׿]{2,40}$', line):
            data["name"] = line
            break

    # Email
    m = re.search(r'[\w.+-]+@[\w-]+\.[a-z]{2,}', text, re.I)
    if m:
        data["email"] = m.group()

    # Phone
    m = re.search(r'0\d[\d\-]{7,10}', text)
    if m:
        data["phone"] = m.group()

    # LinkedIn
    m = re.search(r'linkedin\.com/in/[\w-]+', text, re.I)
    if m:
        data["linkedin"] = "https://" + m.group()

    # Skills — look for common tech keywords
    skill_patterns = [
        "SQL", "Python", "Power BI", "Tableau", "BigQuery", "Snowflake", "dbt",
        "DAX", "Power Query", "Excel", "VBA", "Pandas", "NumPy",
        "Streamlit", "Jupyter", "R", "Looker", "Redshift", "Spark",
        "Jira", "Monday.com", "Figma", "Notion",
        "Machine Learning", "Statistics", "A/B Testing",
        "Data Modeling", "ETL", "KPI", "Dashboard",
    ]
    found_skills = [s for s in skill_patterns if re.search(re.escape(s), text, re.I)]
    if found_skills:
        data["skills"] = found_skills

    # Experience — simple heuristic: find company blocks
    experience = []
    exp_pattern = re.compile(
        r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|'
        r'March|April|June|July|August|September|October|November|December)'
        r'[\s\d]{4,10}|20\d\d)',
        re.I
    )
    date_positions = [m.start() for m in exp_pattern.finditer(text)]
    for i, pos in enumerate(date_positions[:8]):
        snippet = text[max(0, pos - 200):pos + 100].strip()
        if len(snippet) > 20:
            experience.append(snippet[:300])
    if experience:
        data["experience_raw"] = experience

    # Summary — look for "Summary" or "Profile" section
    summary_match = re.search(
        r'(?:summary|profile|about me)[:\n]+([\s\S]{50,400}?)(?:\n{2,}|\Z)',
        text, re.I
    )
    if summary_match:
        data["summary"] = summary_match.group(1).strip()

    return data


def list_cv_files() -> list[str]:
    if not os.path.exists(CV_FOLDER):
        return []
    return [
        f for f in os.listdir(CV_FOLDER)
        if f.lower().endswith((".pdf", ".docx"))
    ]


# ── UI ────────────────────────────────────────────────────────────────────────

def render(lang: str):
    rtl = "rtl" if lang == "he" else "ltr"
    profile = st.session_state.profile.copy()

    title = "👤 הפרופיל שלי" if lang == "he" else "👤 My Profile"
    st.markdown(f'<div class="section-title">{title}</div>', unsafe_allow_html=True)

    tab_labels = (
        ["פרטים אישיים", "ניסיון", "כישורים", "העדפות", "ייבוא קו\"ח"]
        if lang == "he"
        else ["Personal Info", "Experience", "Skills", "Preferences", "Import CV"]
    )
    tabs = st.tabs(tab_labels)

    # ── Tab 0: Personal Info ─────────────────────────────────────────────────
    with tabs[0]:
        c1, c2 = st.columns(2)
        with c1:
            profile["name"] = st.text_input(
                "שם מלא" if lang == "he" else "Full Name",
                value=profile.get("name", ""),
            )
            profile["email"] = st.text_input(
                "אימייל" if lang == "he" else "Email",
                value=profile.get("email", ""),
            )
            profile["phone"] = st.text_input(
                "טלפון" if lang == "he" else "Phone",
                value=profile.get("phone", ""),
            )
        with c2:
            profile["linkedin"] = st.text_input(
                "LinkedIn URL",
                value=profile.get("linkedin", ""),
            )
            profile["location"] = st.text_input(
                "מיקום" if lang == "he" else "Location",
                value=profile.get("location", "Tel Aviv"),
            )
            profile["summary"] = st.text_area(
                "תקציר מקצועי" if lang == "he" else "Professional Summary",
                value=profile.get("summary", ""),
                height=120,
            )

    # ── Tab 1: Experience ────────────────────────────────────────────────────
    with tabs[1]:
        exp_label = "ניסיון תעסוקתי" if lang == "he" else "Work Experience"
        st.markdown(f"**{exp_label}**")
        hint = "הזן כל תפקיד בשורה חדשה: תפקיד | חברה | תאריך התחלה | תאריך סיום | תיאור"
        hint_en = "One role per line: Role | Company | Start | End | Description"
        st.caption(hint if lang == "he" else hint_en)

        default_exp = profile.get("experience_text", "")
        if not default_exp and profile.get("experience_raw"):
            default_exp = "\n\n".join(profile["experience_raw"][:5])

        profile["experience_text"] = st.text_area(
            exp_label,
            value=default_exp,
            height=250,
            label_visibility="collapsed",
        )

        st.markdown("---")
        edu_label = "השכלה" if lang == "he" else "Education"
        st.markdown(f"**{edu_label}**")
        profile["education"] = st.text_area(
            edu_label,
            value=profile.get("education", ""),
            height=100,
            label_visibility="collapsed",
        )

        st.markdown("---")
        mil_label = "שירות צבאי / אחר" if lang == "he" else "Military / Other"
        profile["military"] = st.text_input(
            mil_label,
            value=profile.get("military", ""),
        )

    # ── Tab 2: Skills ────────────────────────────────────────────────────────
    with tabs[2]:
        all_skills_default = [
            "SQL", "Python", "Power BI", "Tableau", "BigQuery", "dbt",
            "DAX", "Power Query", "Excel", "Pandas", "NumPy",
            "Data Modeling", "KPI Design", "Dashboard Development",
            "Statistical Analysis", "A/B Testing", "Stakeholder Management",
            "Project Management", "Process Improvement",
        ]
        current_skills = profile.get("skills", all_skills_default)

        skills_label = "כישורים טכניים ורכים" if lang == "he" else "Technical & Soft Skills"
        hint2 = "הכנס כישורים מופרדים בפסיקים" if lang == "he" else "Comma-separated skills"
        skills_str = st.text_area(
            skills_label,
            value=", ".join(current_skills),
            height=150,
            help=hint2,
        )
        profile["skills"] = [s.strip() for s in skills_str.split(",") if s.strip()]

        cert_label = "הסמכות וקורסים" if lang == "he" else "Certifications & Courses"
        profile["certifications"] = st.text_area(
            cert_label,
            value=profile.get("certifications",
                "The Complete Investment Banking Course\n"
                "SQL for Data Science Bootcamp\n"
                "Microsoft Power BI for Business Intelligence\n"
                "The Data Science Course: Complete Bootcamp"),
            height=120,
        )

        lang_label = "שפות" if lang == "he" else "Languages"
        profile["languages"] = st.text_input(
            lang_label,
            value=profile.get("languages", "Hebrew (Native), English (Fluent)"),
        )

    # ── Tab 3: Preferences ───────────────────────────────────────────────────
    with tabs[3]:
        c1, c2 = st.columns(2)
        with c1:
            pref_roles_label = "תפקידים מועדפים" if lang == "he" else "Target Roles"
            profile["target_roles"] = st.multiselect(
                pref_roles_label,
                options=TARGET_ROLES,
                default=profile.get("target_roles", ["Data Analyst", "Business Analyst"]),
            )

            salary_label = "טווח שכר רצוי (₪ ברוטו/חודש)" if lang == "he" else "Desired Salary Range (₪ gross/month)"
            col_min, col_max = st.columns(2)
            with col_min:
                profile["salary_min"] = st.number_input(
                    "מינימום" if lang == "he" else "Min",
                    value=int(profile.get("salary_min", 18000)),
                    step=1000,
                )
            with col_max:
                profile["salary_max"] = st.number_input(
                    "מקסימום" if lang == "he" else "Max",
                    value=int(profile.get("salary_max", 30000)),
                    step=1000,
                )

        with c2:
            work_type_label = "סוג עבודה" if lang == "he" else "Work Type"
            profile["work_type"] = st.multiselect(
                work_type_label,
                options=["Hybrid", "Remote", "On-site"],
                default=profile.get("work_type", ["Hybrid", "Remote"]),
            )

            priority_label = "מה חשוב לך בתפקיד הבא?" if lang == "he" else "What matters most in your next role?"
            profile["priorities"] = st.text_area(
                priority_label,
                value=profile.get("priorities",
                    "Growth opportunities, interesting data problems, strong team, good culture"),
                height=100,
            )

    # ── Tab 4: Import CV ─────────────────────────────────────────────────────
    with tabs[4]:
        import_title = "ייבוא מקו\"ח קיים" if lang == "he" else "Import from Existing CV"
        st.markdown(f"**{import_title}**")

        cv_files = list_cv_files()
        if cv_files:
            # Show all CV files (PDF + DOCX), grouped by type indicator
            def _file_label(f):
                ext = "📄" if f.lower().endswith(".pdf") else "📝"
                return f"{ext} {f}"

            pick_label = "בחר קובץ לייבוא" if lang == "he" else "Select a CV file to import"
            chosen = st.selectbox(
                pick_label,
                cv_files,
                format_func=_file_label,
            )

            col_btn, col_info = st.columns([1, 2])
            with col_btn:
                btn_label = "📥 " + ("ייבא פרטים" if lang == "he" else "Import details")
                do_import = st.button(btn_label, type="primary", use_container_width=True)
            with col_info:
                ext = chosen.rsplit(".", 1)[-1].upper() if chosen else ""
                st.caption(f"קובץ {ext} — {chosen}" if lang == "he" else f"{ext} file selected")

            if do_import and chosen:
                path = os.path.join(CV_FOLDER, chosen)
                parsed = parse_cv_file(path)
                if parsed:
                    imported = []
                    for k, v in parsed.items():
                        if k not in profile or not profile[k]:
                            profile[k] = v
                            imported.append(k)
                    fields_str = ", ".join(imported) if imported else "—"
                    st.success("✅ " + ("יובאו: " + fields_str if lang == "he"
                                        else "Imported: " + fields_str))
                else:
                    st.warning("לא הצלחנו לחלץ פרטים. נסה קובץ אחר." if lang == "he"
                               else "Could not extract details. Try another file.")
        else:
            st.info("תיקיית קורות החיים לא נמצאה." if lang == "he"
                    else f"CV folder not found at: {CV_FOLDER}")

        st.markdown("---")
        upload_label = "או גרור קובץ חדש" if lang == "he" else "Or upload a new file"
        uploaded = st.file_uploader(upload_label, type=["pdf", "docx"])
        if uploaded:
            import tempfile
            suffix = "." + uploaded.name.rsplit(".", 1)[-1].lower()
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(uploaded.read())
                tmp_path = tmp.name
            parsed = parse_cv_file(tmp_path)
            os.unlink(tmp_path)
            if parsed:
                for k, v in parsed.items():
                    if k not in profile or not profile[k]:
                        profile[k] = v
                st.success("✅ " + ("יובא!" if lang == "he" else "Imported!"))

    # ── Save button ───────────────────────────────────────────────────────────
    st.markdown("<br>", unsafe_allow_html=True)
    save_label = "💾 שמור פרופיל" if lang == "he" else "💾 Save Profile"
    if st.button(save_label, type="primary"):
        save_profile(profile)
        st.success("✅ " + ("פרופיל נשמר!" if lang == "he" else "Profile saved!"))

import io
import os
import re
import json
import html as _html
import streamlit as st
from config import ANTHROPIC_API_KEY, CLAUDE_FAST


# ── CV Templates ──────────────────────────────────────────────────────────────

TEMPLATES = {
    "🎨 Modern Blue": {
        "accent":      "#2563eb",
        "name_color":  "#0f172a",
        "text_color":  "#1e293b",
        "sub_color":   "#64748b",
        "bg":          "#ffffff",
        "border":      "#e2e8f0",
        "font_name":   "Arial, sans-serif",
        "font_body":   "Arial, sans-serif",
        "name_size":   "1.5rem",
        "section_case": "upper",
    },
    "🖤 Executive Dark": {
        "accent":      "#1e293b",
        "name_color":  "#0f172a",
        "text_color":  "#1e293b",
        "sub_color":   "#475569",
        "bg":          "#f8fafc",
        "border":      "#cbd5e1",
        "font_name":   "Georgia, serif",
        "font_body":   "Georgia, serif",
        "name_size":   "1.6rem",
        "section_case": "upper",
    },
    "🌿 Minimal Green": {
        "accent":      "#059669",
        "name_color":  "#065f46",
        "text_color":  "#1e293b",
        "sub_color":   "#6b7280",
        "bg":          "#ffffff",
        "border":      "#d1fae5",
        "font_name":   "Helvetica, sans-serif",
        "font_body":   "Helvetica, sans-serif",
        "name_size":   "1.4rem",
        "section_case": "title",
    },
}


# ── Claude tailoring ──────────────────────────────────────────────────────────

def tailor_cv(profile: dict, job_description: str, language: str = "en") -> dict:
    if not ANTHROPIC_API_KEY or ANTHROPIC_API_KEY == "your_key_here":
        return _fallback_tailor(profile, job_description)

    import anthropic
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

    skills    = ", ".join(profile.get("skills", []))
    exp       = profile.get("experience_text", "")
    summary   = profile.get("summary", "")
    education = profile.get("education", "")

    prompt = f"""You are a senior CV consultant and professional writer.
Tailor this candidate's CV to closely match the job description below.

STRICT RULES:
- Do NOT invent new experience, companies, or skills not present in the original
- Rewrite bullet points to mirror keywords from the job description naturally
- Make the summary punchy and role-specific (2-3 sentences max)
- Return ONLY valid JSON, no markdown, no extra text

JOB DESCRIPTION:
{job_description[:2500]}

CANDIDATE PROFILE:
Summary: {summary}
Experience: {exp[:2000]}
Skills: {skills}
Education: {education}

Return this exact JSON structure:
{{
  "summary": "<tailored 2-3 sentence professional summary>",
  "experience_bullets": ["<bullet 1>", "<bullet 2>", "<bullet 3>", "<bullet 4>", "<bullet 5>"],
  "highlighted_skills": ["<top 10 most relevant skills>"],
  "keywords_used": ["<keyword1>", "<keyword2>"],
  "tailoring_notes": "<1 sentence about what was emphasised>"
}}"""

    try:
        msg = client.messages.create(
            model=CLAUDE_FAST,
            max_tokens=1500,
            messages=[{"role": "user", "content": prompt}],
        )
        text = msg.content[0].text.strip()
        m = re.search(r'\{[\s\S]+\}', text)
        if m:
            result = json.loads(m.group())
            return {**profile, **result, "_tailored": True}
    except Exception:
        pass

    return _fallback_tailor(profile, job_description)


def _fallback_tailor(profile: dict, jd: str) -> dict:
    return {**profile, "_tailored": False}


def get_keywords(jd_text: str) -> list[str]:
    """Extract meaningful tech/role keywords from a job description."""
    if not jd_text:
        return []
    stopwords = {
        # Function words
        "with", "and", "the", "for", "you", "will", "have", "that", "this",
        "are", "our", "your", "from", "able", "been", "must", "should", "would",
        "could", "also", "such", "more", "than", "into", "over", "they", "them",
        "their", "what", "when", "where", "which", "who", "how", "all", "any",
        "both", "each", "other", "about", "through", "various", "within",
        "across", "multiple", "focused", "related", "based", "required",
        # Generic job-desc words (not meaningful as keywords)
        "experience", "skills", "working", "strong", "work", "role", "team",
        "join", "help", "good", "great", "excellent", "years", "using", "build",
        "including", "looking", "make", "provide", "ensure", "well", "high",
        "large", "complex", "player", "economy", "challenges", "environment",
        "opportunities", "impact", "culture", "fast", "lead", "grow", "cross",
        "drive", "world", "time", "need", "best", "know", "take", "like", "real",
        "create", "background", "responsibilities", "requirements", "position",
        "candidate", "ability", "proven", "knowledge", "develop", "maintain",
        "support", "solve", "deliver", "define", "collaborate", "communicate",
        "implement", "design", "analyze", "analyze", "relevant", "understand",
        "ensure", "together", "level", "people", "process", "product", "company",
        "deeply", "hands", "ideal", "major", "minimum", "nice", "overall",
        "passionate", "preferred", "successful", "strong",
    }
    words = re.findall(r'\b[a-zA-Z]{4,}\b', jd_text.lower())
    freq: dict[str, int] = {}
    for w in words:
        if w not in stopwords:
            freq[w] = freq.get(w, 0) + 1
    return [w for w, c in sorted(freq.items(), key=lambda x: -x[1]) if c >= 2][:20]


def ai_extract_keywords(jd_text: str) -> list[str]:
    """Claude-based ATS keyword extraction — hard skills, tools, certifications,
    domain terms. Falls back to [] on failure (caller uses get_keywords then)."""
    if not jd_text.strip():
        return []
    from modules.career_advisor import _claude, _extract_json
    prompt = (
        "Extract the 15-20 most important ATS keywords/phrases from this job "
        "description: hard skills, tools, certifications, methodologies, domain "
        "terms. Skip soft skills and generic words. Return a JSON array of "
        "lowercase strings only.\n\n" + jd_text[:3000]
    )
    raw = _claude(prompt, max_tokens=400, fast=True)
    kws = _extract_json(raw, kind="array")
    return [str(k).lower() for k in kws if isinstance(k, str)][:20]


def improve_bullet(bullet: str, missing_keywords: list[str]) -> list[str]:
    """Rewrite one resume bullet (X-Y-Z formula, strong verb). 2 variants."""
    from modules.career_advisor import _claude, _extract_json
    kw_part = (
        f"If truthful to the original, naturally weave in one of these missing "
        f"job keywords: {', '.join(missing_keywords[:3])}. Never invent facts."
        if missing_keywords else "Never invent facts not present in the original."
    )
    prompt = f"""Rewrite this resume bullet using a strong action verb and the X-Y-Z formula
(Accomplished X, as measured by Y, by doing Z). Keep it one line, max 30 words.
{kw_part}

BULLET: {bullet}

Return a JSON array with exactly 2 rewritten variants (strings)."""
    raw = _claude(prompt, max_tokens=300, fast=True)
    return [str(v) for v in _extract_json(raw, kind="array")][:2]


def ats_check(cv_data: dict, jd_text: str = "", keywords: list[str] | None = None) -> dict:
    warnings = []
    score = 100

    if not cv_data.get("summary"):
        warnings.append("Missing professional summary")
        score -= 10
    elif len(str(cv_data.get("summary", ""))) > 600:
        warnings.append("Summary too long (keep under ~4 lines)")
        score -= 5
    if not cv_data.get("skills"):
        warnings.append("No skills section")
        score -= 15
    elif len(cv_data.get("skills", [])) > 20:
        warnings.append("Too many skills listed (>20) — ATS parsers favor a focused list")
        score -= 5
    if not cv_data.get("experience_text") and not cv_data.get("experience_bullets"):
        warnings.append("No experience section")
        score -= 20
    if not cv_data.get("education"):
        warnings.append("No education listed")
        score -= 5
    if not cv_data.get("email") and not cv_data.get("phone"):
        warnings.append("Missing contact info (email/phone)")
        score -= 10

    bullets = cv_data.get("experience_bullets", [])
    if bullets and not any(re.search(r"\d", b) for b in bullets):
        warnings.append("No quantified results in bullets — add numbers/percentages")
        score -= 5

    keywords = keywords if keywords is not None else get_keywords(jd_text)
    cv_text = (
        str(cv_data.get("summary", "")) + " " +
        str(cv_data.get("experience_text", "")) + " " +
        " ".join(cv_data.get("experience_bullets", [])) + " " +
        " ".join(cv_data.get("skills", []))
    ).lower()

    matched = [k for k in keywords if k in cv_text]
    missing = [k for k in keywords if k not in cv_text][:5]

    if keywords:
        kw_ratio = len(matched) / len(keywords)
        if kw_ratio < 0.3:
            warnings.append(f"Low keyword match ({int(kw_ratio*100)}%)")
            score -= 15
        elif kw_ratio < 0.6:
            score -= 5

    return {
        "score": max(score, 0),
        "warnings": warnings,
        "matched_keywords": matched[:8],
        "missing_keywords": missing,
        "keyword_ratio": int(len(matched) / max(len(keywords), 1) * 100),
    }


# ── DOCX export ───────────────────────────────────────────────────────────────

def export_docx(cv_data: dict, language: str = "en", template_name: str = "🎨 Modern Blue") -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return b""

    tpl = TEMPLATES.get(template_name, TEMPLATES["🎨 Modern Blue"])
    accent_hex = tpl["accent"].lstrip("#")
    accent_rgb = tuple(int(accent_hex[i:i+2], 16) for i in (0, 2, 4))

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    def _rgb(hex_color):
        h = hex_color.lstrip("#")
        return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(cv_data.get("name", ""))
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = _rgb(tpl["name_color"])

    # Contact
    parts = [x for x in [cv_data.get("phone"), cv_data.get("email"),
                          cv_data.get("linkedin"), cv_data.get("location")] if x]
    if parts:
        cp = doc.add_paragraph(" · ".join(parts))
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cp.runs:
            r.font.size = Pt(9)
            r.font.color.rgb = _rgb(tpl["sub_color"])

    _add_divider(doc, accent_rgb)

    def _section(heading, lines, bullets=False):
        label = heading.upper() if tpl["section_case"] == "upper" else heading.title()
        _add_section_header(doc, label, accent_rgb)
        for line in lines:
            if line.strip():
                pp = doc.add_paragraph(("• " if bullets else "") + line.strip(),
                                       style="List Bullet" if bullets else "Normal")
                pp.runs[0].font.size = Pt(10) if pp.runs else None

    if cv_data.get("summary"):
        _section("Professional Summary" if language == "en" else "תקציר מקצועי",
                 [cv_data["summary"]])

    bullets = cv_data.get("experience_bullets", [])
    exp     = cv_data.get("experience_text", "")
    if bullets or exp:
        lines = bullets if bullets else [l for l in exp.splitlines() if l.strip()]
        _section("Experience" if language == "en" else "ניסיון", lines, bullets=bool(bullets))

    if cv_data.get("education"):
        _section("Education" if language == "en" else "השכלה",
                 [l for l in cv_data["education"].splitlines() if l.strip()])

    skills = cv_data.get("highlighted_skills") or cv_data.get("skills", [])
    if skills:
        _section("Skills" if language == "en" else "כישורים", [" · ".join(skills)])

    if cv_data.get("certifications"):
        _section("Certifications" if language == "en" else "הסמכות",
                 [l for l in cv_data["certifications"].splitlines() if l.strip()], bullets=True)

    if cv_data.get("languages"):
        _section("Languages" if language == "en" else "שפות", [cv_data["languages"]])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_section_header(doc, text, accent_rgb):
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(*accent_rgb)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    color_hex = "{:02X}{:02X}{:02X}".format(*accent_rgb)
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1");   bot.set(qn("w:color"), color_hex)
    pBdr.append(bot); pPr.append(pBdr)


def _add_divider(doc, accent_rgb):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    color_hex = "{:02X}{:02X}{:02X}".format(*accent_rgb)
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "1");   bot.set(qn("w:color"), color_hex)
    pBdr.append(bot); pPr.append(pBdr)


# ── PDF export ────────────────────────────────────────────────────────────────

def export_pdf(cv_data: dict, language: str = "en", template_name: str = "🎨 Modern Blue") -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
    except ImportError:
        return b""

    tpl    = TEMPLATES.get(template_name, TEMPLATES["🎨 Modern Blue"])
    ACCENT = colors.HexColor(tpl["accent"])
    DARK   = colors.HexColor(tpl["name_color"])
    TEXT   = colors.HexColor(tpl["text_color"])
    GRAY   = colors.HexColor(tpl["sub_color"])

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=1.5*cm, bottomMargin=1.5*cm)

    name_s    = ParagraphStyle("n", fontSize=20, textColor=DARK,
                                alignment=TA_CENTER, fontName="Helvetica-Bold", spaceAfter=3)
    contact_s = ParagraphStyle("c", fontSize=9,  textColor=GRAY,
                                alignment=TA_CENTER, spaceAfter=6)
    section_s = ParagraphStyle("s", fontSize=11, textColor=ACCENT,
                                fontName="Helvetica-Bold", spaceBefore=10, spaceAfter=2)
    body_s    = ParagraphStyle("b", fontSize=10, textColor=TEXT, leading=14, spaceAfter=2)
    bullet_s  = ParagraphStyle("bl", fontSize=10, textColor=TEXT, leading=14,
                                leftIndent=12, spaceAfter=2)

    story = [Paragraph(cv_data.get("name", ""), name_s)]

    parts = [x for x in [cv_data.get("phone"), cv_data.get("email"),
                          cv_data.get("location"), cv_data.get("linkedin")] if x]
    story.append(Paragraph(" · ".join(parts), contact_s))
    story.append(HRFlowable(width="100%", thickness=1.5, color=ACCENT))

    def add_section(heading, lines, is_bullets=False):
        label = heading.upper() if tpl["section_case"] == "upper" else heading.title()
        story.append(Paragraph(label, section_s))
        story.append(HRFlowable(width="100%", thickness=0.5, color=ACCENT))
        story.append(Spacer(1, 2))
        for line in lines:
            if line.strip():
                style = bullet_s if is_bullets else body_s
                story.append(Paragraph(("• " if is_bullets else "") + line.strip(), style))

    if cv_data.get("summary"):
        add_section("Professional Summary" if language == "en" else "תקציר מקצועי",
                    [cv_data["summary"]])

    bullets = cv_data.get("experience_bullets", [])
    exp     = cv_data.get("experience_text", "")
    if bullets or exp:
        lines = bullets if bullets else [l for l in exp.splitlines() if l.strip()]
        add_section("Experience" if language == "en" else "ניסיון", lines, is_bullets=bool(bullets))

    if cv_data.get("education"):
        add_section("Education" if language == "en" else "השכלה",
                    [l for l in cv_data["education"].splitlines() if l.strip()])

    skills = cv_data.get("highlighted_skills") or cv_data.get("skills", [])
    if skills:
        add_section("Skills" if language == "en" else "כישורים", [" · ".join(skills)])

    if cv_data.get("certifications"):
        add_section("Certifications" if language == "en" else "הסמכות",
                    [l for l in cv_data["certifications"].splitlines() if l.strip()], is_bullets=True)

    if cv_data.get("languages"):
        add_section("Languages" if language == "en" else "שפות", [cv_data["languages"]])

    doc.build(story)
    return buf.getvalue()


# ── UI ─────────────────────────────────────────────────────────────────────────

def render(lang: str):
    profile      = st.session_state.profile
    selected_job = st.session_state.get("selected_job")

    st.markdown(
        '<div class="section-title">📄 ' +
        ("בניית קו\"ח" if lang == "he" else "CV Builder") + "</div>",
        unsafe_allow_html=True,
    )

    if not profile or not profile.get("name"):
        st.warning("⚠️ " + ("אנא השלם את הפרופיל תחילה." if lang == "he"
                             else "Please complete your profile first."))
        if st.button("👤 " + ("עבור לפרופיל" if lang == "he" else "Go to Profile")):
            st.session_state.page = "profile"
            st.rerun()
        return

    if "cv_jd_text" not in st.session_state:
        st.session_state.cv_jd_text = ""
    # Always use fresh profile as base; only AI tailoring should produce highlighted_skills
    if "tailored_cv" not in st.session_state:
        st.session_state.tailored_cv = profile.copy()
    else:
        # Sync base fields (name, email, skills, education, etc.) from current profile
        # so stale AI-tailored skills don't override real profile on fresh page load
        base = st.session_state.tailored_cv
        if not base.get("_tailored"):
            st.session_state.tailored_cv = profile.copy()

    # When a new job is navigated to from Job Search, sync its description into the JD field
    _cur_job_id = (selected_job or {}).get("id", "")
    if selected_job and _cur_job_id != st.session_state.get("_cv_last_job_id", ""):
        st.session_state.cv_jd_text = selected_job.get("description", "") or ""
        st.session_state._cv_last_job_id = _cur_job_id
        st.session_state.tailored_cv = profile.copy()  # reset to base profile for new job

    # ── Top bar: template + language ──────────────────────────────────────────
    col_tpl, col_lang, col_spacer = st.columns([2, 1, 2])
    with col_tpl:
        template_name = st.selectbox(
            "עיצוב" if lang == "he" else "Template",
            list(TEMPLATES.keys()),
            label_visibility="collapsed",
        )
    with col_lang:
        cv_lang_label = st.radio(
            "שפה" if lang == "he" else "Lang",
            ["EN", "HE"],
            horizontal=True,
            label_visibility="collapsed",
        )
        cv_lang_code = "he" if cv_lang_label == "HE" else "en"

    st.markdown('<div style="height:0.5rem"></div>', unsafe_allow_html=True)

    # ── Main: two columns ─────────────────────────────────────────────────────
    col_left, col_right = st.columns([1, 1], gap="large")

    with col_left:
        # Job description
        st.markdown(
            '<p style="font-weight:600;color:#e6edf3;margin-bottom:0.3rem">' +
            ("📋 תיאור המשרה" if lang == "he" else "📋 Job Description") + "</p>",
            unsafe_allow_html=True,
        )

        if selected_job:
            jd_missing = not selected_job.get("description", "").strip()
            notice_color = "#f59e0b" if jd_missing else "#2563eb"
            notice_border = "#f59e0b" if jd_missing else "#2563eb"
            notice_icon = "⚠️" if jd_missing else "✅"
            st.markdown(
                f'<div style="background:#1a2744;border:1px solid {notice_border};border-radius:6px;'
                f'padding:0.5rem 0.75rem;margin-bottom:0.5rem;font-size:0.85rem">'
                f'<span style="color:{notice_color}">{notice_icon} {_html.escape(selected_job["title"])}</span>'
                f' <span style="color:#6b7280">@ {_html.escape(selected_job["company"])}</span>'
                + (f'<br><span style="color:#9ca3af;font-size:0.78rem">LinkedIn jobs don\'t include full description — paste it below</span>' if jd_missing else "")
                + f"</div>",
                unsafe_allow_html=True,
            )

        jd_input = st.text_area(
            "jd",
            value=st.session_state.cv_jd_text,
            height=260,
            placeholder="הדבק תיאור משרה כאן..." if lang == "he" else "Paste job description here...",
            label_visibility="collapsed",
            key="cv_jd_input",
        )
        st.session_state.cv_jd_text = jd_input

        btn_col, reset_col = st.columns([4, 1])
        with btn_col:
            do_tailor = st.button(
                "✨ " + ("התאם קו\"ח למשרה עם AI" if lang == "he" else "Tailor CV with AI"),
                type="primary",
                use_container_width=True,
            )
        with reset_col:
            if st.button("🔄", use_container_width=True,
                         help="Reset to base profile" if lang == "en" else "איפוס לפרופיל המקורי"):
                st.session_state.tailored_cv = profile.copy()
                st.session_state.pop("_tailored_flag", None)
                st.rerun()
        if do_tailor:
            if not jd_input.strip():
                st.error("נא להזין תיאור משרה" if lang == "he" else "Please enter a job description first.")
            else:
                with st.spinner("✨ " + ("Claude מתאים את קו\"ח שלך..." if lang == "he"
                                         else "Claude is tailoring your CV...")):
                    tailored = tailor_cv(profile, jd_input, cv_lang_code)
                    st.session_state.tailored_cv = tailored
                if tailored.get("_tailored"):
                    if tailored.get("tailoring_notes"):
                        st.success(f"✅ {tailored['tailoring_notes']}")
                    else:
                        st.success("✅ " + ("קו\"ח הותאם בהצלחה!" if lang == "he" else "CV tailored successfully!"))
                else:
                    st.info("ℹ️ " + ("Claude לא זמין — מציג קו\"ח בסיסי" if lang == "he"
                                     else "Claude not available — showing base CV"))

        # ATS + keywords
        st.markdown('<div style="height:0.75rem"></div>', unsafe_allow_html=True)
        has_claude = bool(ANTHROPIC_API_KEY) and ANTHROPIC_API_KEY != "your_key_here"
        jd_hash = str(hash(jd_input.strip()))
        ai_kw_cache = st.session_state.setdefault("ai_kw_cache", {})

        if st.button("🤖 " + ("מילות מפתח AI" if lang == "he" else "AI keywords"),
                     disabled=not has_claude or not jd_input.strip(),
                     help=("חילוץ מילות מפתח מדויק עם Claude במקום ספירת מילים"
                           if lang == "he" else
                           "Precise keyword extraction with Claude instead of word counting")):
            with st.spinner("מחלץ מילות מפתח..." if lang == "he" else "Extracting keywords..."):
                kws = ai_extract_keywords(jd_input)
                if kws:
                    ai_kw_cache[jd_hash] = kws

        ai_keywords = ai_kw_cache.get(jd_hash)  # None → frequency fallback
        ats = ats_check(st.session_state.tailored_cv, jd_input, keywords=ai_keywords)
        _render_ats_panel(ats, lang)
        _render_improve_bullet(st.session_state.tailored_cv, ats, lang, has_claude)

    with col_right:
        st.markdown(
            '<p style="font-weight:600;color:#e6edf3;margin-bottom:0.3rem">' +
            ("👁️ תצוגה מקדימה" if lang == "he" else "👁️ Live Preview") + "</p>",
            unsafe_allow_html=True,
        )
        _render_cv_preview(st.session_state.tailored_cv, cv_lang_code, template_name)

    # ── Export bar ────────────────────────────────────────────────────────────
    st.markdown("---")
    exp_col1, exp_col2, exp_col3 = st.columns([1, 1, 2])

    cv_data  = st.session_state.tailored_cv
    filename = re.sub(r'\s+', '_', cv_data.get("name", "CV"))

    with exp_col1:
        docx_bytes = export_docx(cv_data, cv_lang_code, template_name)
        if docx_bytes:
            st.download_button(
                "📥 הורד DOCX" if lang == "he" else "📥 Download DOCX",
                data=docx_bytes,
                file_name=f"{filename}_CV.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
    with exp_col2:
        pdf_bytes = export_pdf(cv_data, cv_lang_code, template_name)
        if pdf_bytes:
            st.download_button(
                "📥 הורד PDF" if lang == "he" else "📥 Download PDF",
                data=pdf_bytes,
                file_name=f"{filename}_CV.pdf",
                mime="application/pdf",
                use_container_width=True,
            )

    # ── Inline editor ─────────────────────────────────────────────────────────
    st.markdown('<div style="height:0.5rem"></div>', unsafe_allow_html=True)
    with st.expander("✏️ " + ("עריכה ידנית" if lang == "he" else "Manual Edit"), expanded=False):
        cv = st.session_state.tailored_cv

        he = lang == "he"
        tab_sum, tab_exp, tab_skills, tab_extra = st.tabs([
            "תקציר" if he else "Summary",
            "ניסיון" if he else "Experience",
            "כישורים" if he else "Skills",
            "השכלה ועוד" if he else "Education & More",
        ])
        with tab_sum:
            cv["summary"] = st.text_area(
                "תקציר מקצועי" if he else "Professional Summary",
                value=cv.get("summary", ""), height=120)

        with tab_exp:
            bullets = cv.get("experience_bullets", [])
            if bullets:
                st.caption("שורות שנוצרו ע\"י AI (שורה לכל bullet):" if he
                           else "AI-generated bullets (one per line):")
                bullets_text = st.text_area(
                    "שורות ניסיון" if he else "Experience bullets",
                    value="\n".join(bullets), height=200)
                cv["experience_bullets"] = [b.strip() for b in bullets_text.splitlines() if b.strip()]
            cv["experience_text"] = st.text_area(
                "טקסט ניסיון מלא" if he else "Full experience text",
                value=cv.get("experience_text", ""), height=150)

        with tab_skills:
            skills_str = ", ".join(cv.get("highlighted_skills") or cv.get("skills", []))
            new_skills = st.text_input(
                "כישורים (מופרדים בפסיק)" if he else "Skills (comma-separated)",
                value=skills_str)
            skill_list = [s.strip() for s in new_skills.split(",") if s.strip()]
            cv["highlighted_skills"] = skill_list
            cv["skills"] = skill_list

        with tab_extra:
            cv["education"]      = st.text_area("השכלה" if he else "Education",
                                                value=cv.get("education", ""),      height=80)
            cv["certifications"] = st.text_area("הסמכות" if he else "Certifications",
                                                value=cv.get("certifications", ""), height=80)
            cv["languages"]      = st.text_input("שפות" if he else "Languages",
                                                 value=cv.get("languages", ""))

        st.session_state.tailored_cv = cv


# ── CV Preview helpers ────────────────────────────────────────────────────────

def _build_exp_html(bullets: list, exp_raw: str) -> str:
    """Render experience section. AI bullets → flat list. Raw text → parsed structure."""
    if bullets:
        return "".join(
            f'<div style="display:flex;gap:5px;margin-bottom:3px">'
            f'<span style="flex-shrink:0;margin-top:1px">•</span>'
            f'<span>{_html.escape(b)}</span></div>'
            for b in bullets
        )
    if not exp_raw:
        return '<span style="color:#9ca3af;font-style:italic">No experience added yet</span>'

    html_parts = []
    lines = exp_raw.splitlines()
    prev_was_title = False

    for line in lines:
        s = line.strip()
        if not s:
            prev_was_title = False
            continue

        has_sep  = '•' in s
        has_year = bool(re.search(r'\b\d{4}\b', s))

        # Company • Location • Dates line: italic gray — follows a job title
        if prev_was_title and (has_sep or has_year):
            html_parts.append(
                f'<div style="font-style:italic;color:#475569;'
                f'font-size:0.76rem;margin-bottom:5px">{_html.escape(s)}</div>'
            )
            prev_was_title = False
            continue

        # Bullet line
        if s[0] in ('•', '-', '▸', '*', '–'):
            clean = s[1:].strip()
            html_parts.append(
                f'<div style="display:flex;gap:6px;margin-bottom:3px;padding-left:2px">'
                f'<span style="flex-shrink:0;margin-top:1px;color:#374151">•</span>'
                f'<span>{_html.escape(clean)}</span></div>'
            )
            prev_was_title = False
            continue

        # Job title: no bullet char, no • separator, short, starts uppercase, no year
        if not has_sep and not has_year and len(s) < 70 and s[0].isupper():
            spacer = '<div style="height:0.6rem"></div>' if html_parts else ''
            html_parts.append(
                f'{spacer}<div style="font-weight:700;font-size:0.82rem;'
                f'margin-bottom:2px">{_html.escape(s)}</div>'
            )
            prev_was_title = True
            continue

        # Fallback: plain bullet
        html_parts.append(
            f'<div style="display:flex;gap:5px;margin-bottom:2px">'
            f'<span style="flex-shrink:0;margin-top:1px">•</span>'
            f'<span>{_html.escape(s)}</span></div>'
        )
        prev_was_title = False

    return ''.join(html_parts)


def _build_edu_html(edu: str) -> str:
    if not edu:
        return ""
    parts = []
    for line in edu.splitlines():
        s = line.strip()
        if not s:
            continue
        if '•' in s:
            degree, _, rest = s.partition('•')
            parts.append(
                f'<div style="margin-bottom:3px">'
                f'<strong>{_html.escape(degree.strip())}</strong>'
                f' &nbsp;•&nbsp; {_html.escape(rest.strip())}</div>'
            )
        else:
            parts.append(f'<div style="margin-bottom:3px">{_html.escape(s)}</div>')
    return ''.join(parts)


# ── CV Preview ────────────────────────────────────────────────────────────────

def _render_cv_preview(cv: dict, lang: str, template_name: str = "🎨 Modern Blue"):
    tpl    = TEMPLATES.get(template_name, TEMPLATES["🎨 Modern Blue"])
    accent = tpl["accent"]
    rtl    = "rtl" if lang == "he" else "ltr"

    name    = _html.escape(cv.get("name", ""))
    summary = _html.escape(cv.get("summary", ""))
    langs   = _html.escape(cv.get("languages", ""))

    raw_loc      = cv.get("location", "")
    raw_phone    = cv.get("phone", "")
    raw_email    = cv.get("email", "")
    raw_linkedin = cv.get("linkedin", "")

    _a = 'color:{sub};text-decoration:none'.format(sub=tpl["sub_color"])
    contact_items = []
    if raw_loc:
        contact_items.append(_html.escape(raw_loc))
    if raw_phone:
        contact_items.append(
            f'<a href="tel:{_html.escape(raw_phone)}" style="{_a}">'
            f'{_html.escape(raw_phone)}</a>'
        )
    if raw_email:
        contact_items.append(
            f'<a href="mailto:{_html.escape(raw_email)}" style="{_a}">'
            f'{_html.escape(raw_email)}</a>'
        )
    if raw_linkedin:
        li_url = raw_linkedin if raw_linkedin.startswith("http") else f"https://{raw_linkedin}"
        contact_items.append(
            f'<a href="{_html.escape(li_url)}" target="_blank" style="{_a};font-weight:500">LinkedIn</a>'
        )
    contact_html = " &nbsp;·&nbsp; ".join(contact_items)

    skills      = cv.get("highlighted_skills") or cv.get("skills", [])
    skills_html = " &nbsp;·&nbsp; ".join(_html.escape(s) for s in skills[:20]) if skills else ""

    exp_html = _build_exp_html(cv.get("experience_bullets", []), cv.get("experience_text", ""))
    edu_html = _build_edu_html(cv.get("education", ""))

    certifications = cv.get("certifications", "")
    certs_html = "".join(
        f'<div style="display:flex;gap:5px;margin-bottom:2px">'
        f'<span style="flex-shrink:0">•</span>'
        f'<span>{_html.escape(l)}</span></div>'
        for l in certifications.splitlines() if l.strip()
    ) if certifications else ""

    section_label = lambda t: (t.upper() if tpl["section_case"] == "upper" else t.title())

    def sec_hdr(title):
        return (
            f'<div style="color:{accent};font-weight:700;font-size:0.72rem;'
            f'letter-spacing:0.08em;border-bottom:1.5px solid {accent};'
            f'padding-bottom:2px;margin:0.85rem 0 0.45rem 0">{section_label(title)}</div>'
        )

    def section(title, body_html):
        if not body_html:
            return ""
        return (
            sec_hdr(title) +
            f'<div style="font-size:0.78rem;line-height:1.6;'
            f'color:{tpl["text_color"]};margin-bottom:0.1rem">{body_html}</div>'
        )

    has_exp = bool(cv.get("experience_bullets") or cv.get("experience_text"))

    parts = [
        f'<div dir="{rtl}" style="background:{tpl["bg"]};color:{tpl["text_color"]};'
        f'padding:1.4rem 1.8rem 1.8rem;border-radius:6px;border:1px solid {tpl["border"]};'
        f'font-family:{tpl["font_body"]};font-size:0.8rem;line-height:1.5;min-height:480px">',

        # Centered name + contact
        '<div style="text-align:center;padding-bottom:0.6rem;border-bottom:none">',
        f'<div style="font-family:{tpl["font_name"]};font-size:{tpl["name_size"]};'
        f'font-weight:800;color:{tpl["name_color"]};letter-spacing:0.04em;'
        f'line-height:1.2;margin-bottom:0.35rem">{name}</div>',
        f'<div style="color:{tpl["sub_color"]};font-size:0.75rem;line-height:1.6;'
        f'word-break:break-word">{contact_html}</div>',
        '</div>',

        section("Professional Summary" if lang == "en" else "תקציר מקצועי",
                f"<div>{summary}</div>") if summary else "",

        section("Experience" if lang == "en" else "ניסיון", exp_html) if has_exp else "",

        section("Education" if lang == "en" else "השכלה", edu_html) if edu_html else "",

        section("Skills" if lang == "en" else "כישורים",
                f'<div>{skills_html}</div>') if skills_html else "",

        section("Certifications" if lang == "en" else "הסמכות", certs_html) if certs_html else "",

        section("Languages" if lang == "en" else "שפות",
                f'<div>{langs}</div>') if langs else "",

        '</div>',
    ]

    html = "".join(parts)
    st.markdown(html, unsafe_allow_html=True)


# ── ATS Panel ─────────────────────────────────────────────────────────────────

def _render_ats_panel(ats: dict, lang: str):
    score = ats["score"]
    color = "#22c55e" if score >= 80 else ("#f59e0b" if score >= 60 else "#ef4444")
    kw_ratio = ats.get("keyword_ratio", 0)

    st.markdown(f"""
    <div style="background:#161b22;border:1px solid #30363d;border-radius:8px;
                padding:0.9rem 1rem;margin-bottom:0.5rem">
      <div style="display:flex;align-items:center;gap:1rem;margin-bottom:0.6rem">
        <div>
          <div style="font-size:0.7rem;color:#6b7280;text-transform:uppercase;
                      letter-spacing:0.06em">ATS Score</div>
          <div style="font-size:1.8rem;font-weight:800;color:{color};line-height:1">
            {score}<span style="font-size:0.9rem">%</span>
          </div>
        </div>
        <div style="flex:1">
          <div style="background:#21262d;border-radius:4px;height:8px;overflow:hidden">
            <div style="background:{color};width:{score}%;height:100%;
                        border-radius:4px;transition:width 0.3s"></div>
          </div>
          <div style="font-size:0.7rem;color:#6b7280;margin-top:4px">
            {"Keyword match" if lang == "en" else "התאמת מילות מפתח"}: {kw_ratio}%
          </div>
        </div>
      </div>
    """, unsafe_allow_html=True)

    # Warnings
    for w in ats.get("warnings", []):
        st.markdown(
            f'<div style="font-size:0.75rem;color:#f59e0b;margin-bottom:3px">⚠️ {_html.escape(w)}</div>',
            unsafe_allow_html=True,
        )

    # Matched keywords
    matched = ats.get("matched_keywords", [])
    missing = ats.get("missing_keywords", [])

    kw_html = ""
    for k in matched:
        kw_html += (f'<span style="background:#16423240;color:#4ade80;border:1px solid #22c55e40;'
                    f'border-radius:4px;padding:1px 6px;font-size:0.7rem;margin:2px">'
                    f'✓ {_html.escape(k)}</span>')
    for k in missing:
        kw_html += (f'<span style="background:#7f1d1d20;color:#f87171;border:1px solid #ef444440;'
                    f'border-radius:4px;padding:1px 6px;font-size:0.7rem;margin:2px">'
                    f'✗ {_html.escape(k)}</span>')
    # Single close for the panel div opened above — one code path, no dangling tag.
    st.markdown(
        f'<div style="margin-top:0.3rem;line-height:2">{kw_html}</div></div>',
        unsafe_allow_html=True,
    )


def _render_improve_bullet(cv_data: dict, ats: dict, lang: str, has_claude: bool):
    """'Improve this bullet' action — X-Y-Z rewrite with missing-keyword weaving."""
    bullets = cv_data.get("experience_bullets") or [
        b.strip().lstrip("•-· ") for b in str(cv_data.get("experience_text", "")).splitlines()
        if b.strip() and len(b.strip()) > 15
    ]
    if not bullets:
        return

    with st.expander("✨ " + ("שפר ניסוח bullet" if lang == "he" else "Improve a bullet")):
        chosen = st.selectbox(
            "בחר שורה" if lang == "he" else "Pick a line",
            bullets, key="improve_bullet_select",
            format_func=lambda b: b[:80] + ("…" if len(b) > 80 else ""),
        )
        if st.button("✨ " + ("שכתב" if lang == "he" else "Rewrite"),
                     key="improve_bullet_btn", disabled=not has_claude):
            with st.spinner("משכתב..." if lang == "he" else "Rewriting..."):
                variants = improve_bullet(chosen, ats.get("missing_keywords", []))
            if variants:
                st.session_state["bullet_variants"] = (chosen, variants)
            else:
                st.warning("⚠️ " + ("נסה שוב" if lang == "he" else "Try again"))

        cached = st.session_state.get("bullet_variants")
        if cached and cached[0] == chosen:
            for v in cached[1]:
                st.code(v, language=None)
